"""Document ingest tests for the solo-owner on-ramp."""

from __future__ import annotations

import csv
import re
import tomllib
from datetime import UTC, date, datetime
from io import BytesIO, StringIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from complyos.connectors.base import LMSConnector
from complyos.connectors.capabilities import get_connector_capability
from complyos.connectors.document import DocumentExtractionError, DocumentExtractor
from complyos.core.repository import LocalRepository
from complyos.models.domain import (
    AuditReport,
    Course,
    LearningRecord,
    LearningRecordStatus,
    User,
)
from complyos.services.context import default_local_context
from complyos.services.imports import ImportPreviewRequest, ImportService
from complyos.web.dashboard import create_dashboard_app

HEADERS = [
    "Learner ID",
    "Learner Email",
    "Learner Name",
    "Training ID",
    "Training",
    "Status",
    "Completed Date",
    "Renewal Date",
    "Score",
]
ROW = [
    "u1",
    "alex@example.com",
    "Alex Rivera",
    "fall-protection",
    "Fall Protection",
    "completed",
    "2025-01-15",
    "2027-01-15",
    "97",
]


class EmptyAuditor:
    async def generate_report(
        self,
        department: str | None = None,
        region: str | None = None,
    ) -> AuditReport:
        return AuditReport(
            generated_at=datetime(2026, 1, 1, tzinfo=UTC),
            scope="all",
            total_users_audited=0,
            gaps_found=0,
            gaps_by_severity={},
            gaps_by_department={},
            top_missing_courses=[],
            evidence_hash="document-ingest",
            details=[],
        )


def _write_docx(path: Path, rows: list[list[str]]) -> Path:
    document = Document()
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            table.cell(row_index, column_index).text = value
    document.save(path)
    return path


def _docx_bytes_without_table() -> bytes:
    document = Document()
    document.add_paragraph("No primary table in this document.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _write_xlsx(path: Path, rows: list[list[str]]) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    workbook.save(path)
    return path


def _write_csv(path: Path, rows: list[list[str]]) -> Path:
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)
    return path


@pytest.mark.parametrize(
    ("suffix", "writer"),
    [
        (".docx", _write_docx),
        (".xlsx", _write_xlsx),
        (".csv", _write_csv),
    ],
)
async def test_document_extractor_round_trips_primary_table_to_connector_models(
    tmp_path: Path,
    suffix: str,
    writer,
) -> None:
    path = writer(tmp_path / f"records{suffix}", [HEADERS, ROW])

    connector = DocumentExtractor(path)

    assert isinstance(connector, LMSConnector)
    assert await connector.authenticate() is True
    assert await connector.trigger_reminder("u1", "fall-protection") is False

    users = await connector.get_users()
    courses = await connector.get_courses()
    enrollments = await connector.get_enrollments()
    records = await connector.get_learning_records()

    assert [user.id for user in users] == ["u1"]
    assert users[0].email == "alex@example.com"
    assert users[0].first_name == "Alex"
    assert users[0].last_name == "Rivera"
    assert [course.id for course in courses] == ["fall-protection"]
    assert courses[0].title == "Fall Protection"
    assert enrollments[0].user_id == "u1"
    assert enrollments[0].course_id == "fall-protection"
    assert records[0].status == LearningRecordStatus.COMPLETED
    assert records[0].completed_date == datetime(2025, 1, 15)
    assert records[0].expires_at == date(2027, 1, 15)
    assert records[0].score == 97.0
    assert records[0].raw_source_hash is not None

    import_csv = connector.to_import_csv_text()
    assert "user_id,course_id,status" in import_csv
    assert "u1,fall-protection,completed" in import_csv


def test_document_extractor_supports_uploaded_bytes(tmp_path: Path) -> None:
    path = _write_xlsx(tmp_path / "records.xlsx", [HEADERS, ROW])

    connector = DocumentExtractor(path.read_bytes(), filename="records.xlsx")

    assert connector.to_import_csv_text().splitlines()[1].startswith(
        "document-upload-0,u1,fall-protection"
    )


def test_document_extractor_fails_closed_for_missing_or_unsupported_primary_table(
    tmp_path: Path,
) -> None:
    empty_docx = tmp_path / "empty.docx"
    empty_docx.write_bytes(_docx_bytes_without_table())
    pdf = tmp_path / "records.pdf"
    pdf.write_bytes(b"%PDF-1.7")

    with pytest.raises(DocumentExtractionError, match="primary table"):
        DocumentExtractor(empty_docx).to_import_csv_text()

    with pytest.raises(DocumentExtractionError, match="unsupported file type"):
        DocumentExtractor(pdf).to_import_csv_text()


def test_document_upload_table_goes_through_preview_decide_promote_governance(
    tmp_path: Path,
) -> None:
    governance_headers = [*HEADERS, "Source Record ID"]
    table_rows = [
        governance_headers,
        [*ROW, "dupe-1"],
        [
            "u1",
            "alex@example.com",
            "Alex Rivera",
            "fall-protection",
            "Fall Protection",
            "completed",
            "2025-01-16",
            "2027-01-15",
            "98",
            "dupe-1",
        ],
    ]
    path = _write_csv(tmp_path / "dupes.csv", table_rows)
    import_csv = DocumentExtractor(path).to_import_csv_text()
    repo = LocalRepository(str(tmp_path / "document-import.db"))
    service = ImportService(repo)
    context = default_local_context(surface="shell", role="import_approver")

    preview = service.preview(
        context,
        ImportPreviewRequest(source_system="document_upload", csv_text=import_csv),
    )
    blocked = service.promote(context, preview.batch_id)
    records_after_blocked_promote = repo.list_learning_records(source_system="document_upload")
    rows = repo.list_import_rows(preview.batch_id)
    needs_decision = next(row for row in rows if row["validation_status"] == "NEEDS_DECISION")
    service.decide(
        context,
        batch_id=preview.batch_id,
        row_id=needs_decision["id"],
        decision_type="accept",
        reason="duplicate reviewed",
    )
    promoted = service.promote(context, preview.batch_id)

    assert preview.status == "QUARANTINED"
    assert preview.row_counts["VALID"] == 1
    assert preview.row_counts["NEEDS_DECISION"] == 1
    assert preview.can_promote is False
    assert blocked.status == "QUARANTINED"
    assert records_after_blocked_promote == []
    assert promoted.status == "PROMOTED"
    assert promoted.promoted_rows == 2
    assert len(repo.list_learning_records(source_system="document_upload")) == 2
    assert repo.list_evidence_ledger()[0]["query_type"] == "import.promote"
    assert {"import.preview", "import.decide", "import.promote"} <= {
        item["action"] for item in repo.list_action_logs(limit=20)
    }


def _local_client(monkeypatch: pytest.MonkeyPatch, tmp_path: Path, db_name: str):
    monkeypatch.delenv("COMPLYOS_API_TOKEN", raising=False)
    monkeypatch.setenv("COMPLYOS_ALLOW_INSECURE_LOCAL", "1")
    repo = LocalRepository(str(tmp_path / db_name))
    app = create_dashboard_app(auditor=EmptyAuditor(), repository=repo)
    return TestClient(app), repo


def _login_local(client: TestClient, role: str) -> None:
    client.post("/shell/login", data={"role": role}, follow_redirects=False)


def _seed_record(repo: LocalRepository) -> None:
    repo.save_user(
        User(
            id="u1",
            employee_id="u1",
            email="alex@example.com",
            first_name="Alex",
            last_name="Rivera",
            department="Field",
            region="US",
            hire_date=date(2024, 1, 1),
            custom_attributes={"tenant_id": "local-default"},
        )
    )
    repo.save_course(
        Course(
            id="fall-protection",
            code="FALL-1",
            title="Fall Protection",
            mandatory=True,
        )
    )
    repo.save_learning_record(
        LearningRecord(
            id="lr-expired",
            user_id="u1",
            course_id="fall-protection",
            source_system="document_upload",
            status=LearningRecordStatus.COMPLETED,
            completed_date=datetime(2025, 1, 15),
            expires_at=date(2025, 6, 1),
        )
    )


def test_shell_document_upload_missing_table_surfaces_preview_issue(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    client, _repo = _local_client(monkeypatch, tmp_path, "shell-doc-upload.db")
    _login_local(client, "importer")

    response = client.post(
        "/shell/imports/preview",
        data={"csv_text": ""},
        files={
            "document_file": (
                "empty.docx",
                _docx_bytes_without_table(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert "PARTIAL_LOAD" in response.text
    assert "No" in response.text


def test_shell_records_view_and_status_packet_exports_use_expiry_status(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    client, repo = _local_client(monkeypatch, tmp_path, "shell-records.db")
    _seed_record(repo)
    _login_local(client, "compliance_manager")

    view = client.get("/shell/records")
    csv_export = client.get("/shell/records/export.csv")
    html_export = client.get("/shell/records/export.html")

    assert view.status_code == 200
    assert "Alex Rivera" in view.text
    assert "Fall Protection" in view.text
    assert "2025-01-15" in view.text
    assert "2025-06-01" in view.text
    assert "overdue" in view.text
    assert csv_export.status_code == 200
    assert "learner,training,completed_date,renewal_date,status" in csv_export.text
    assert "Alex Rivera,Fall Protection,2025-01-15,2025-06-01,overdue" in csv_export.text
    assert html_export.status_code == 200
    assert "client-facing status packet" in html_export.text
    assert "overdue" in html_export.text


def test_document_ingest_new_surfaces_stay_inside_claim_language(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    client, repo = _local_client(monkeypatch, tmp_path, "shell-record-claims.db")
    _seed_record(repo)
    _login_local(client, "compliance_manager")
    surfaces = [
        client.get("/shell/records").text,
        client.get("/shell/records/export.csv").text,
        client.get("/shell/records/export.html").text,
    ]

    forbidden = re.compile(r"\b(LMS|compliant|certified|audit-proof|replaces)\b", re.I)
    for text in surfaces:
        assert forbidden.search(text) is None


def test_document_upload_connector_capability_is_registered() -> None:
    capability = get_connector_capability("document_upload")

    assert capability.display_name == "Document Upload"
    assert capability.profile == "both"
    assert capability.status == "supported"
    assert capability.auth == "none"
    assert capability.supports_learning_records is True
    assert capability.supports_expiry is True


def test_document_ingest_dependencies_are_declared_without_pdf_extraction() -> None:
    pyproject = tomllib.loads(Path("pyproject.toml").read_text(encoding="utf-8"))
    dependencies = pyproject["project"]["dependencies"]

    assert any(dependency.startswith("python-docx") for dependency in dependencies)
    assert any(dependency.startswith("openpyxl") for dependency in dependencies)
    assert not any(dependency.startswith("pdfplumber") for dependency in dependencies)


def test_status_packet_csv_has_exact_columns(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    client, repo = _local_client(monkeypatch, tmp_path, "shell-record-columns.db")
    _seed_record(repo)
    _login_local(client, "compliance_manager")

    response = client.get("/shell/records/export.csv")
    reader = csv.DictReader(StringIO(response.text))

    assert reader.fieldnames == [
        "learner",
        "training",
        "completed_date",
        "renewal_date",
        "status",
    ]
